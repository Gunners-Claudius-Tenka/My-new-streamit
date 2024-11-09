#先写一下这个计划吧
#(1.摄像头) 2.读取文件(dcm) (3.生成docx文件(不确定能不能行)) 4.图像增强 5.考虑一下布局美观 6.上传本地照片
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from datetime import datetime
import io
import tempfile
import os

st.markdown("<h1 style='font-family: 宋体; font-size: 24px; text-align: center;'>这是一个关于streamlit和opencv的练习网站</h1>", unsafe_allow_html=True)
name_input = st.text_input("您的名字是：","")
st.write("确认是否输入完成：",name_input)
picture = st.camera_input("Take a picture")
if picture:
    # 将图像数据转换为PIL图像
    image = Image.open(picture)
    # 创建一个临时文件
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    # 保存图像到临时文件
    image.save(temp_file.name, format='PNG')
    temp_file.close()  # 关闭文件

    # 获取时间
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    # 创建Word文档
    doc = Document()
    heading = doc.add_heading('这是一个关于照片的测试', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run('当前时间：').bold = True  # 将“当前时间：”设置为粗体
    p.add_run(current_time)  # 添加当前时间
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #添加一个用于储存输入的段落
    p = doc.add_paragraph()
    p.add_run("您输入的名字是：")
    p.add_run(name_input)
    # 添加一个居中的段落用于图片
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(temp_file.name, width=Inches(4))  # 设置图片宽度并居中

    # 保存Word文档到临时文件
    temp_doc_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_doc_file.name)
    temp_doc_file.close()  # 关闭文件

    # 下载按钮
    st.download_button(
        label="Download Word Document",
        data=open(temp_doc_file.name, 'rb').read(),
        file_name='picture.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    # 删除临时文件
    os.remove(temp_file.name)
    os.remove(temp_doc_file.name)